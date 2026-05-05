import io
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib.patches import Polygon as MplPolygon
from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BRAND_BLUE = RGBColor(0x1F, 0x4E, 0x79)
OK_GREEN = RGBColor(0x2E, 0x7D, 0x32)
FAIL_RED = RGBColor(0xC6, 0x28, 0x28)
MUTED_GRAY = RGBColor(0x66, 0x66, 0x66)
COL_FACE_RGBA = (1.0, 0.54, 0.40, 0.16)
COL_EDGE = '#D84315'
COL_TEXT = '#BF360C'


def _shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tc_pr.append(shd)
    shd.set(qn('w:fill'), fill)


def _apply_report_style(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Arial')
    normal.font.size = Pt(9.5)

    for style_name, size, color in [
        ('Title', 18, BRAND_BLUE),
        ('Heading 1', 13, BRAND_BLUE),
        ('Heading 2', 10.5, BRAND_BLUE),
    ]:
        style = doc.styles[style_name]
        style.font.name = 'Arial'
        style._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Arial')
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run('STM Pile Cap Design Calculation | ACI 318-19')
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED_GRAY

def _norm_col(c):
    if isinstance(c, dict):
        return (c.get("section", "Square"),
                c.get("bx", 500.0), c.get("by", 500.0),
                c.get("diam", 500.0))
    return ("Square", float(c), float(c), float(c))


def _col_pos(c):
    if isinstance(c, dict):
        return float(c.get("x", 0.0)), float(c.get("y", 0.0))
    return 0.0, 0.0

def _format_col(c):
    sec, bx, by, dm = _norm_col(c)
    x, y = _col_pos(c)
    if sec == "Circular":
        base = "Circular, D = {:.0f} mm".format(dm)
    elif sec == "Rectangular":
        base = "Rectangular, {:.0f} x {:.0f} mm".format(bx, by)
    else:
        base = "Square, {:.0f} mm side".format(bx)
    if abs(x) > 1e-6 or abs(y) > 1e-6:
        return "{} at ({:.0f}, {:.0f}) mm".format(base, x, y)
    return base

def _norm_pile(p):
    if isinstance(p, dict):
        return (p.get("section", "Circular"),
                p.get("bx", 600.0), p.get("by", 600.0),
                p.get("diam", 600.0))
    return ("Circular", float(p), float(p), float(p))

def _format_pile(p):
    sec, bx, by, dm = _norm_pile(p)
    if sec == "Circular":
        return "Circular, D = {:.0f} mm".format(dm)
    if sec == "Square":
        return "Square, {:.0f} mm side".format(bx)
    return "Rectangular, {:.0f} x {:.0f} mm".format(bx, by)


def _rebar_spacing(n_bars, distribution_width_mm, cap_lx_mm, cap_ly_mm,
                   cover_mm):
    n = int(n_bars)
    if n <= 1:
        return None
    edge_inset = min(max(float(cover_mm), 25.0),
                     0.45 * min(float(cap_lx_mm), float(cap_ly_mm)))
    usable_width = max(0.0, float(distribution_width_mm) - 2.0 * edge_inset)
    return usable_width / (n - 1)


def _format_spacing(value):
    if value is None:
        return "-"
    return "{:.0f}".format(value)

def _plot_plan(coords, D, lx, ly, cx, cy, col_size, cap_polygon, results):
    fig, ax = plt.subplots(figsize=(7, 7))
    if cap_polygon:
        ax.add_patch(MplPolygon(cap_polygon, closed=True, fill=True,
            facecolor='#bdc3c7', edgecolor='#34495e',
            linewidth=2, alpha=0.4))
    else:
        ax.add_patch(patches.Rectangle(
            (cx-lx/2, cy-ly/2), lx, ly,
            facecolor='#bdc3c7', edgecolor='#34495e',
            linewidth=2, alpha=0.4))
    sec_c, cbx, cby, cdm = _norm_col(col_size)
    col_x, col_y = _col_pos(col_size)
    if sec_c == "Circular":
        ax.add_patch(patches.Circle(
            (col_x, col_y), cdm/2,
            facecolor=COL_FACE_RGBA, edgecolor=COL_EDGE, linewidth=2))
    else:
        ax.add_patch(patches.Rectangle(
            (col_x-cbx/2, col_y-cby/2), cbx, cby,
            facecolor=COL_FACE_RGBA, edgecolor=COL_EDGE, linewidth=2))
    ax.text(col_x, col_y, 'COL', ha='center', va='center',
            color=COL_TEXT, fontsize=9, fontweight='bold')
    pile_loads = results.get("pile_loads_kN", [])
    sec_p, pbx, pby, pdm = _norm_pile(D)
    for i, (x, y) in enumerate(coords, 1):
        if sec_p == "Circular":
            ax.add_patch(patches.Circle((x, y), pdm/2,
                facecolor='#5B8DEF', edgecolor='#1F4E89', linewidth=1.5))
        elif sec_p == "Square":
            ax.add_patch(patches.Rectangle(
                (x-pbx/2, y-pbx/2), pbx, pbx,
                facecolor='#5B8DEF', edgecolor='#1F4E89', linewidth=1.5))
        else:
            ax.add_patch(patches.Rectangle(
                (x-pbx/2, y-pby/2), pbx, pby,
                facecolor='#5B8DEF', edgecolor='#1F4E89', linewidth=1.5))
        lbl = 'P{}'.format(i)
        if i-1 < len(pile_loads):
            lbl += '\n{:.0f}kN'.format(pile_loads[i-1])
        ax.text(x, y, lbl, ha='center', va='center',
                color='white', fontsize=8, fontweight='bold')
        ax.plot([col_x, x], [col_y, y], 'r--', linewidth=1.5, alpha=0.7)
    pad = max(lx, ly)*0.15 + 300
    ax.set_xlim(min(cx-lx/2, col_x)-pad, max(cx+lx/2, col_x)+pad)
    ax.set_ylim(min(cy-ly/2, col_y)-pad, max(cy+ly/2, col_y)+pad)
    ax.set_aspect('equal')
    ax.set_title('Plan View - Pile Cap Layout & Struts')
    ax.set_xlabel('X (mm)')
    ax.set_ylabel('Y (mm)')
    ax.grid(True, alpha=0.3)
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=120, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _plot_bottom_rebar_fig(coords, D, lx, ly, cx, cy, col_size,
                           cap_polygon, x_bar, x_n, y_bar, y_n,
                           x_chk, y_chk, cover_mm=75):
    """Matplotlib plan view of bottom-face reinforcement for report Section 5."""
    _DB = {"DB12": 113.10, "DB16": 201.06, "DB20": 314.16,
           "DB25": 490.87, "DB28": 615.75, "DB32": 804.25}

    fig, ax = plt.subplots(figsize=(7.5, 7.5))
    # Cap outline
    if cap_polygon:
        ax.add_patch(MplPolygon(cap_polygon, closed=True, fill=True,
            facecolor='#ecf0f1', edgecolor='#2c3e50', linewidth=2, alpha=0.5))
    else:
        ax.add_patch(patches.Rectangle(
            (cx-lx/2, cy-ly/2), lx, ly,
            facecolor='#ecf0f1', edgecolor='#2c3e50', linewidth=2, alpha=0.5))
    # Column
    sec_c, cbx, cby, cdm = _norm_col(col_size)
    col_x, col_y = _col_pos(col_size)
    if sec_c == "Circular":
        ax.add_patch(patches.Circle((col_x, col_y), cdm/2,
            facecolor=COL_FACE_RGBA, edgecolor=COL_EDGE, linewidth=2, zorder=2))
    else:
        ax.add_patch(patches.Rectangle(
            (col_x-cbx/2, col_y-cby/2), cbx, cby,
            facecolor=COL_FACE_RGBA, edgecolor=COL_EDGE, linewidth=2, zorder=2))
    # Piles
    sec_p, pbx, pby, pdm = _norm_pile(D)
    hw_x = (pdm if sec_p == "Circular" else pbx) / 2.0
    hw_y = (pdm if sec_p == "Circular" else pby) / 2.0
    for i, (px, py) in enumerate(coords, 1):
        if sec_p == "Circular":
            ax.add_patch(patches.Circle((px, py), pdm/2,
                facecolor='#5B8DEF', edgecolor='#1F4E89', linewidth=1.5, zorder=4))
        else:
            ax.add_patch(patches.Rectangle(
                (px-hw_x, py-hw_y), 2*hw_x, 2*hw_y,
                facecolor='#5B8DEF', edgecolor='#1F4E89', linewidth=1.5, zorder=4))
        ax.text(px, py, 'P{}'.format(i), ha='center', va='center',
                color='white', fontsize=8, fontweight='bold', zorder=5)

    if cap_polygon:
        poly = list(cap_polygon)
        xs_cap = [p[0] for p in poly]
        ys_cap = [p[1] for p in poly]
        cap_x_min, cap_x_max = min(xs_cap), max(xs_cap)
        cap_y_min, cap_y_max = min(ys_cap), max(ys_cap)
    else:
        poly = [
            (cx - lx/2, cy - ly/2),
            (cx + lx/2, cy - ly/2),
            (cx + lx/2, cy + ly/2),
            (cx - lx/2, cy + ly/2),
        ]
        cap_x_min, cap_x_max = cx - lx/2, cx + lx/2
        cap_y_min, cap_y_max = cy - ly/2, cy + ly/2

    edge_inset = min(max(float(cover_mm), 25.0), 0.45 * min(lx, ly))

    def _linspace(a, b, n):
        if n <= 1 or b <= a:
            return [(a + b) / 2.0]
        return [a + (b - a) * k / (n - 1) for k in range(n)]

    def _horizontal_segment(y):
        hits = []
        for k, (x1, y1) in enumerate(poly):
            x2, y2 = poly[(k + 1) % len(poly)]
            if abs(y2 - y1) < 1e-9:
                continue
            if (y >= min(y1, y2)) and (y < max(y1, y2)):
                t = (y - y1) / (y2 - y1)
                hits.append(x1 + t * (x2 - x1))
        hits.sort()
        if len(hits) >= 2:
            return hits[0] + edge_inset, hits[-1] - edge_inset
        return cap_x_min + edge_inset, cap_x_max - edge_inset

    def _vertical_segment(x):
        hits = []
        for k, (x1, y1) in enumerate(poly):
            x2, y2 = poly[(k + 1) % len(poly)]
            if abs(x2 - x1) < 1e-9:
                continue
            if (x >= min(x1, x2)) and (x < max(x1, x2)):
                t = (x - x1) / (x2 - x1)
                hits.append(y1 + t * (y2 - y1))
        hits.sort()
        if len(hits) >= 2:
            return hits[0] + edge_inset, hits[-1] - edge_inset
        return cap_y_min + edge_inset, cap_y_max - edge_inset

    x_min = cap_x_min + edge_inset; x_max = cap_x_max - edge_inset
    y_min = cap_y_min + edge_inset; y_max = cap_y_max - edge_inset

    # Bottom X-bars (horizontal, red solid)
    ys_bar = _linspace(y_min, y_max, x_n)
    for yb in ys_bar:
        x0, x1 = _horizontal_segment(yb)
        if x1 <= x0:
            continue
        ax.plot([x0, x1], [yb, yb], color='#E91E63', linewidth=2.0,
                solid_capstyle='round', zorder=3)

    # Bottom Y-bars (vertical, blue dashed)
    xs_bar = _linspace(x_min, x_max, y_n)
    for xb in xs_bar:
        y0, y1 = _vertical_segment(xb)
        if y1 <= y0:
            continue
        ax.plot([xb, xb], [y0, y1], color='#3F51B5', linewidth=2.0,
                linestyle='--', dashes=(8, 4), zorder=3)

    # Legend proxies
    ax.plot([], [], color='#E91E63', linewidth=2, label='{} × {} (X-dir) As={:.0f}mm²'.format(
        x_n, x_bar, x_chk.get('As_provided', x_n*_DB.get(x_bar, 0))))
    ax.plot([], [], color='#3F51B5', linewidth=2, linestyle='--',
            label='{} × {} (Y-dir) As={:.0f}mm²'.format(
        y_n, y_bar, y_chk.get('As_provided', y_n*_DB.get(y_bar, 0))))

    # Status annotations
    sx = "OK" if x_chk.get("ok") else "FAIL"
    sy = "OK" if y_chk.get("ok") else "FAIL"
    ax.annotate(
        'X: As_req={:.0f}  As_prov={:.0f}  → {}'.format(
            x_chk.get('As_required', 0), x_chk.get('As_provided', 0), sx),
        xy=(cx, y_max+300), ha='center', fontsize=9, color='#C62828' if sx=='FAIL' else '#1B5E20')
    ax.annotate(
        'Y: As_req={:.0f}  As_prov={:.0f}  → {}'.format(
            y_chk.get('As_required', 0), y_chk.get('As_provided', 0), sy),
        xy=(cx, y_max+500), ha='center', fontsize=9, color='#C62828' if sy=='FAIL' else '#1B5E20')

    pad = max(lx, ly)*0.1 + 600
    ax.set_xlim(min(cx-lx/2, col_x)-pad, max(cx+lx/2, col_x)+pad)
    ax.set_ylim(min(cy-ly/2, col_y)-pad-100,
                max(cy+ly/2, col_y)+pad+600)
    ax.set_aspect('equal')
    ax.legend(loc='lower center', fontsize=9, framealpha=0.85)
    ax.set_title('Bottom-Face Reinforcement Layout (Plan View)')
    ax.set_xlabel('X (mm)'); ax.set_ylabel('Y (mm)')
    ax.grid(True, alpha=0.25)
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=120, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _plot_top_rebar_fig(coords, D, lx, ly, cx, cy, col_size,
                        cap_polygon, top_rebar, cover_mm=75):
    """Matplotlib plan view of top-face reinforcement for report Section 6."""
    import math as _m
    _DB = {"DB12": 113.10, "DB16": 201.06, "DB20": 314.16,
           "DB25": 490.87, "DB28": 615.75, "DB32": 804.25}

    bar   = top_rebar.get("top_bar_size", "DB20")
    A_bar = _DB.get(bar, 314.16)
    fy_d  = top_rebar.get("fy_design_mpa", 390.0)
    As_x  = top_rebar.get("As_top_x_mm2", 0.0)
    As_y  = top_rebar.get("As_top_y_mm2", 0.0)
    s_max = top_rebar.get("s_max_top_mm", 450.0)
    n_x_area = max(2, int(_m.ceil(As_x / A_bar))) if As_x > 0 else 2
    n_y_area = max(2, int(_m.ceil(As_y / A_bar))) if As_y > 0 else 2
    n_x = int(top_rebar.get("top_x_n_bars", n_x_area))
    n_y = int(top_rebar.get("top_y_n_bars", n_y_area))

    fig, ax = plt.subplots(figsize=(7.5, 7.5))
    # Cap outline
    if cap_polygon:
        ax.add_patch(MplPolygon(cap_polygon, closed=True, fill=True,
            facecolor='#fef9e7', edgecolor='#2c3e50', linewidth=2, alpha=0.6))
    else:
        ax.add_patch(patches.Rectangle(
            (cx-lx/2, cy-ly/2), lx, ly,
            facecolor='#fef9e7', edgecolor='#2c3e50', linewidth=2, alpha=0.6))
    # Column
    sec_c, cbx, cby, cdm = _norm_col(col_size)
    col_x, col_y = _col_pos(col_size)
    if sec_c == "Circular":
        ax.add_patch(patches.Circle((col_x, col_y), cdm/2,
            facecolor=COL_FACE_RGBA, edgecolor=COL_EDGE, linewidth=2, zorder=2))
    else:
        ax.add_patch(patches.Rectangle(
            (col_x-cbx/2, col_y-cby/2), cbx, cby,
            facecolor=COL_FACE_RGBA, edgecolor=COL_EDGE, linewidth=2, zorder=2))
    # Piles (outline only — top view)
    sec_p, pbx, pby, pdm = _norm_pile(D)
    hw_x = (pdm if sec_p == "Circular" else pbx) / 2.0
    hw_y = (pdm if sec_p == "Circular" else pby) / 2.0
    for i, (px, py) in enumerate(coords, 1):
        if sec_p == "Circular":
            ax.add_patch(patches.Circle((px, py), pdm/2,
                facecolor='none', edgecolor='#7F8C8D', linewidth=1.5,
                linestyle=':', zorder=3))
        else:
            ax.add_patch(patches.Rectangle(
                (px-hw_x, py-hw_y), 2*hw_x, 2*hw_y,
                facecolor='none', edgecolor='#7F8C8D', linewidth=1.5,
                linestyle=':', zorder=3))
        ax.text(px, py, 'P{}'.format(i), ha='center', va='center',
                color='#5D6D7E', fontsize=8, zorder=4)

    if cap_polygon:
        poly = list(cap_polygon)
        xs_cap = [p[0] for p in poly]
        ys_cap = [p[1] for p in poly]
        cap_x_min, cap_x_max = min(xs_cap), max(xs_cap)
        cap_y_min, cap_y_max = min(ys_cap), max(ys_cap)
    else:
        poly = [
            (cx - lx/2, cy - ly/2),
            (cx + lx/2, cy - ly/2),
            (cx + lx/2, cy + ly/2),
            (cx - lx/2, cy + ly/2),
        ]
        cap_x_min, cap_x_max = cx - lx/2, cx + lx/2
        cap_y_min, cap_y_max = cy - ly/2, cy + ly/2

    edge_inset = min(max(float(cover_mm), 20.0), 0.45 * min(lx, ly))

    def _linspace(a, b, n):
        if n <= 1 or b <= a:
            return [(a + b) / 2.0]
        return [a + (b - a) * k / (n - 1) for k in range(n)]

    def _horizontal_segment(y):
        hits = []
        for k, (x1, y1) in enumerate(poly):
            x2, y2 = poly[(k + 1) % len(poly)]
            if abs(y2 - y1) < 1e-9:
                continue
            if (y >= min(y1, y2)) and (y < max(y1, y2)):
                t = (y - y1) / (y2 - y1)
                hits.append(x1 + t * (x2 - x1))
        hits.sort()
        if len(hits) >= 2:
            return hits[0] + edge_inset, hits[-1] - edge_inset
        return cap_x_min + edge_inset, cap_x_max - edge_inset

    def _vertical_segment(x):
        hits = []
        for k, (x1, y1) in enumerate(poly):
            x2, y2 = poly[(k + 1) % len(poly)]
            if abs(x2 - x1) < 1e-9:
                continue
            if (x >= min(x1, x2)) and (x < max(x1, x2)):
                t = (x - x1) / (x2 - x1)
                hits.append(y1 + t * (y2 - y1))
        hits.sort()
        if len(hits) >= 2:
            return hits[0] + edge_inset, hits[-1] - edge_inset
        return cap_y_min + edge_inset, cap_y_max - edge_inset

    x_min = cap_x_min + edge_inset; x_max = cap_x_max - edge_inset
    y_min = cap_y_min + edge_inset; y_max = cap_y_max - edge_inset

    # Top X-bars (horizontal, amber)
    ys_bar = _linspace(y_min, y_max, n_x)
    for yb in ys_bar:
        x0, x1 = _horizontal_segment(yb)
        if x1 <= x0:
            continue
        ax.plot([x0, x1], [yb, yb], color='#FF8F00', linewidth=2.5,
                solid_capstyle='round', zorder=3)

    # Top Y-bars (vertical, teal dashed)
    xs_bar = _linspace(x_min, x_max, n_y)
    for xb in xs_bar:
        y0, y1 = _vertical_segment(xb)
        if y1 <= y0:
            continue
        ax.plot([xb, xb], [y0, y1], color='#00897B', linewidth=2.5,
                linestyle='--', dashes=(8, 4), zorder=3)

    # s_max reference bar
    ax.annotate('', xy=(x_min+s_max, y_max+200), xytext=(x_min, y_max+200),
                arrowprops=dict(arrowstyle='<->', color='#795548', lw=1.5))
    ax.text(x_min+s_max/2, y_max+350, 's_max={:.0f} mm'.format(s_max),
            ha='center', fontsize=9, color='#795548')

    # Legend and annotations
    ax.plot([], [], color='#FF8F00', linewidth=2.5,
            label='{} × {} (Top X) As={:.0f}mm² fy={:.0f}MPa'.format(
                n_x, bar, n_x*A_bar, fy_d))
    ax.plot([], [], color='#00897B', linewidth=2.5, linestyle='--',
            label='{} × {} (Top Y) As={:.0f}mm² fy={:.0f}MPa'.format(
                n_y, bar, n_y*A_bar, fy_d))

    ax.annotate(
        'Top X: As_req={:.0f}  As_prov={:.0f} mm²  ({})'.format(
            As_x, n_x*A_bar, top_rebar.get('governs_x','')),
        xy=(cx, y_max+600), ha='center', fontsize=9, color='#E65100')
    ax.annotate(
        'Top Y: As_req={:.0f}  As_prov={:.0f} mm²  ({})'.format(
            As_y, n_y*A_bar, top_rebar.get('governs_y','')),
        xy=(cx, y_max+800), ha='center', fontsize=9, color='#00695C')
    if fy_d > 420:
        ax.annotate(
            '⚠ fy={:.0f} MPa → spacing limited by §24.3.2'.format(fy_d),
            xy=(cx, y_max+1000), ha='center', fontsize=9,
            color='#B71C1C',
            bbox=dict(boxstyle='round,pad=0.3', facecolor='#FFCCBC', alpha=0.8))

    pad = max(lx, ly)*0.1 + 600
    ax.set_xlim(min(cx-lx/2, col_x)-pad, max(cx+lx/2, col_x)+pad)
    ax.set_ylim(min(cy-ly/2, col_y)-pad-100,
                max(cy+ly/2, col_y)+pad+1100)
    ax.set_aspect('equal')
    ax.legend(loc='lower center', fontsize=9, framealpha=0.85)
    ax.set_title('Top-Face Minimum Reinforcement Layout (Plan View)')
    ax.set_xlabel('X (mm)'); ax.set_ylabel('Y (mm)')
    ax.grid(True, alpha=0.25)
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=120, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _plot_elev(coords, h_cap, D, col_size, results):
    fig, ax = plt.subplots(figsize=(8, 5))
    if not coords:
        plt.close(fig); return None
    xs = [c[0] for c in coords]
    sec_p, pbx, pby, pdm = _norm_pile(D)
    pwid = pdm if sec_p == "Circular" else pbx
    col_x, _col_y = _col_pos(col_size)
    x_min = min(min(xs)-pwid, col_x-pwid)
    x_max = max(max(xs)+pwid, col_x+pwid)
    ax.add_patch(patches.Rectangle(
        (x_min-200, 0), x_max-x_min+400, h_cap,
        facecolor='#bdc3c7', edgecolor='#34495e',
        linewidth=2, alpha=0.4))
    sec_c, cbx, cby, cdm = _norm_col(col_size)
    cwid = cdm if sec_c == "Circular" else cbx
    ax.add_patch(patches.Rectangle(
        (col_x-cwid/2, h_cap), cwid, 400,
        facecolor='#FF8A65', edgecolor='#D84315', linewidth=2))
    for x, _ in coords:
        ax.add_patch(patches.Rectangle(
            (x-pwid/2, -600), pwid, 600,
            facecolor='#5B8DEF', edgecolor='#1F4E89', linewidth=2))
        ax.plot([col_x, x], [h_cap, pwid/2], 'r-', linewidth=3, alpha=0.8)
    if len(coords) >= 2:
        ax.plot([min(xs), max(xs)], [pwid/2, pwid/2], 'g-', linewidth=4)
        ax.text((min(xs)+max(xs))/2, pwid/2-150,
                'Tie F={:.0f}kN'.format(results['F_tie_max_kN']),
                ha='center', color='green', fontweight='bold')
    ax.set_aspect('equal')
    ax.set_title('Elevation - Struts (red) & Tie (green)')
    ax.set_xlabel('X (mm)'); ax.set_ylabel('Z (mm)')
    ax.grid(True, alpha=0.3)
    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=120, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)
    return buf


def _set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.bold = bold
    r.font.name = 'Arial'
    r._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Arial')
    r.font.size = Pt(8.5)
    if str(text).strip().upper() == "OK":
        r.font.color.rgb = OK_GREEN
        r.bold = True
    elif str(text).strip().upper() == "FAIL":
        r.font.color.rgb = FAIL_RED
        r.bold = True
    elif str(text).strip().upper() in ("DESIGN OK", "DESIGN FAILS"):
        r.font.color.rgb = OK_GREEN if "OK" in str(text).upper() else FAIL_RED
        r.bold = True


def _make_table(doc, headers, rows):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Light Grid Accent 1'
    for j, h in enumerate(headers):
        _set_cell(tbl.rows[0].cells[j], h, bold=True)
        _shade_cell(tbl.rows[0].cells[j], 'D9EAF7')
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            _set_cell(tbl.rows[i+1].cells[j], v)
    return tbl


def _fmt_ok(ok):
    return "OK" if ok else "FAIL"


def generate_report(inputs, results, x_chk, y_chk, pairs,
                    anch_x=None, anch_y=None, opt_x=None, opt_y=None,
                    top_rebar=None):
    doc = Document()
    _apply_report_style(doc)

    # Title
    h = doc.add_heading('STM Pile Cap Design Calculation', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Strut-and-Tie Method | ACI 318-19 | Design Calculation Note')
    r.italic = True
    r.font.color.rgb = MUTED_GRAY

    anch_ok = True
    if anch_x is not None:
        anch_ok = anch_ok and anch_x.get('ok', True)
    if anch_y is not None:
        anch_ok = anch_ok and anch_y.get('ok', True)
    design_ok = bool(results.get('overall_OK') and x_chk.get('ok') and
                     y_chk.get('ok') and anch_ok)

    # 0. Executive Summary
    doc.add_heading('0. Executive Summary', level=1)
    _make_table(doc, ['Item', 'Result'], [
        ['Design status', 'DESIGN OK' if design_ok else 'DESIGN FAILS'],
        ['Pile arrangement', '{} piles | {}'.format(
            results.get('n_piles', len(inputs.get('coords', []))),
            _format_pile(inputs['D']))],
        ['Pile cap geometry', '{} | Lx x Ly x h = {:.0f} x {:.0f} x {:.0f} mm'.format(
            inputs.get('shape_label', 'Pile cap'),
            inputs['cap_lx'], inputs['cap_ly'], inputs['h_cap'])],
        ['Column', _format_col(inputs['col_size'])],
        ['Factored load for reactions', 'Pu_total = {:.1f} kN'.format(
            inputs.get('Pu_total_kN', results.get('Pu_total_kN', inputs['Pu'])))],
        ['Critical concrete checks',
         'Strut DCR {:.2f}, pile bearing DCR {:.2f}, column DCR {:.2f}'.format(
             results.get('strut_DCR', 0.0),
             results.get('bearing_DCR', 0.0),
             results.get('column_DCR', 0.0))],
        ['Bottom reinforcement X',
         '{}-{} | As_prov {:.0f} / As_req {:.0f} mm² | {}'.format(
             x_chk['n_bars'], x_chk['bar_size'],
             x_chk['As_provided'], x_chk['As_required'],
             results.get('As_x_governs', 'governing'))],
        ['Bottom reinforcement Y',
         '{}-{} | As_prov {:.0f} / As_req {:.0f} mm² | {}'.format(
             y_chk['n_bars'], y_chk['bar_size'],
             y_chk['As_provided'], y_chk['As_required'],
             results.get('As_y_governs', 'governing'))],
        ['Top reinforcement basis',
         '(0.0018Ag)/2 each direction' if top_rebar else 'Not evaluated'],
        ['Anchorage',
         _fmt_ok(anch_ok) if (anch_x is not None or anch_y is not None)
         else 'Not evaluated'],
    ])
    p = doc.add_paragraph()
    p.add_run('Model scope: ').bold = True
    p.add_run(
        'STM concrete capacity checks are preliminary and should be reviewed '
        'with final strut/nodal-zone geometry, detailing, constructability, '
        'and project-specific code requirements.')

    # 1. Inputs
    doc.add_heading('1. Design Inputs', level=1)
    doc.add_heading('1.1 Materials', level=2)
    _make_table(doc, ['Parameter', 'Value'], [
        ["Concrete strength f'c", "{} MPa".format(inputs['fc'])],
        ["Steel yield fy", "{} MPa".format(inputs['fy'])],
        ["Concrete cover", "{} mm".format(inputs['cover'])],
    ])

    doc.add_heading('1.2 Geometry', level=2)
    _make_table(doc, ['Parameter', 'Value'], [
        ["Pile section", _format_pile(inputs['D'])],
        ["Cap thickness h", "{} mm".format(inputs['h_cap'])],
        ["Column section", _format_col(inputs['col_size'])],
        ["Number of piles", "{}".format(results['n_piles'])],
        ["Cap shape", inputs['shape_label']],
        ["Cap Lx x Ly", "{:.0f} x {:.0f} mm".format(
            inputs['cap_lx'], inputs['cap_ly'])],
    ])

    doc.add_heading('1.3 Column Loads', level=2)
    _w_cap_nom = inputs.get('W_cap_nom_kN', inputs.get('W_cap_kN', 0.0))
    _uls_fac   = inputs.get('wcap_uls_factor', 1.2)
    _w_cap_uls = inputs.get('W_cap_kN', 0.0)
    _pu_total  = inputs.get('Pu_total_kN', inputs['Pu'])
    _make_table(doc, ['Load', 'Value'], [
        ["Axial Pu (column load, ULS)", "{:.1f} kN".format(inputs['Pu'])],
        ["W_cap nominal  (Lx × Ly × h × 24 kN/m³)",
         "{:.1f} kN".format(_w_cap_nom)],
        ["ULS factor γ for W_cap", "{:.2f}".format(_uls_fac)],
        ["W_cap (ULS) = W_cap_nom × γ",
         "{:.1f} kN".format(_w_cap_uls)],
        ["Pu_total = Pu + W_cap(ULS)  ← used for pile reactions",
         "{:.1f} kN".format(_pu_total)],
        ["Moment Mux (about X-axis)", "{:.1f} kN·m".format(inputs['Mux'])],
        ["Moment Muy (about Y-axis)", "{:.1f} kN·m".format(inputs['Muy'])],
    ])

    # 2. Pile Layout
    doc.add_heading('2. Pile Layout & Reactions', level=1)
    _sec_p, _pbx, _pby, _pdm = _norm_pile(inputs['D'])
    if _sec_p == "Circular":
        _gov_min = _pdm
    elif _sec_p == "Square":
        _gov_min = _pbx
    else:
        _gov_min = min(_pbx, _pby)
    doc.add_paragraph(
        "Per ACI 318-19 §13.4.2.1, minimum pile center-to-center spacing "
        "is 2.5 x governing dim = {:.0f} mm.".format(2.5*_gov_min))

    img = _plot_plan(inputs['coords'], inputs['D'],
                     inputs['cap_lx'], inputs['cap_ly'],
                     inputs['cap_cx'], inputs['cap_cy'],
                     inputs['col_size'], inputs['cap_polygon'], results)
    doc.add_picture(img, width=Inches(5.5))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run('Figure 1: Plan view of pile cap with struts').italic = True

    doc.add_heading('2.1 Pile-to-Pile Spacing Check', level=2)
    p = doc.add_paragraph(
        "Distance computed by: d_ij = sqrt((x_i - x_j)^2 + (y_i - y_j)^2)")
    rows = [["P{}-P{}".format(i, j), "{:.0f}".format(d),
             "{:.0f}".format(clear),
             "OK" if ctc_ok else "FAIL",
             "OK" if clear_ok else "FAIL",
             "OK" if ok else "FAIL"]
            for (i, j, d, clear, ctc_ok, clear_ok, ok) in pairs]
    _make_table(doc, ['Pair', 'Distance (mm)', 'Clear (mm)',
                      'c/c', 'Clear', 'Status'], rows)

    doc.add_heading('2.2 Pile Reactions (Rigid-Cap Elastic)', level=2)
    p = doc.add_paragraph()
    p.add_run('Formula: ').bold = True
    p.add_run("P_i = Pu_total/n + (Mux,c × y'_i) / Σy'_j² + "
              "(Muy,c × x'_i) / Σx'_j²")
    p = doc.add_paragraph()
    p.add_run('where Pu_total = Pu + W_cap = {:.1f} + {:.1f} = {:.1f} kN'.format(
        inputs['Pu'],
        inputs.get('W_cap_kN', 0.0),
        inputs.get('Pu_total_kN', inputs['Pu'])))
    _cgx, _cgy = results.get('pile_group_centroid', (0.0, 0.0))
    p = doc.add_paragraph()
    p.add_run(
        "Pile group centroid = ({:.0f}, {:.0f}) mm. Coordinates x', y' "
        "are measured about this centroid; Mux,c and Muy,c include column "
        "load eccentricity for custom layouts.".format(_cgx, _cgy))
    rows = [["P{}".format(i+1),
             "{:.0f}".format(c[0]),
             "{:.0f}".format(c[1]),
             "{:.1f}".format(P)]
            for i, (c, P) in enumerate(
                zip(inputs['coords'], results['pile_loads_kN']))]
    _make_table(doc, ['Pile', 'X (mm)', 'Y (mm)', 'P_i (kN)'], rows)
    if results['has_uplift']:
        p = doc.add_paragraph()
        rr = p.add_run("Warning: Uplift detected (P_min = {:.1f} kN). "
                       "Provide tension piles or anchorage.".format(
                           results['P_min_kN']))
        rr.font.color.rgb = RGBColor(0xC6, 0x28, 0x28)
        rr.bold = True

    # 3. STM Analysis
    doc.add_heading('3. Strut-and-Tie Analysis', level=1)
    p = doc.add_paragraph()
    p.add_run('Effective depth: ').bold = True
    _sp2, _pbx2, _pby2, _pdm2 = _norm_pile(inputs['D'])
    if _sp2 == "Circular":
        _gov_max = _pdm2
    elif _sp2 == "Square":
        _gov_max = _pbx2
    else:
        _gov_max = max(_pbx2, _pby2)
    p.add_run("d_eff = h - cover - db/2 = {} - {} - 12.5 = {:.0f} mm "
              "(ACI 318-19 §23.2, db assumed = DB25)".format(
                  inputs['h_cap'], inputs['cover'],
                  results['d_effective_mm']))

    img2 = _plot_elev(inputs['coords'], inputs['h_cap'],
                      inputs['D'], inputs['col_size'], results)
    if img2:
        doc.add_picture(img2, width=Inches(6))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run('Figure 2: Elevation showing struts and bottom tie').italic = True

    doc.add_heading('3.1 Strut Forces', level=2)
    p = doc.add_paragraph()
    p.add_run("Each strut force: F_strut_i = P_i x L_strut / d_eff "
              "(equilibrium at column-pile node)")
    rows = []
    for i, s in enumerate(results['struts']):
        rows.append([
            "S{}".format(i+1),
            "{:.0f}".format(s['L_strut']),
            "{:.1f}".format(s['theta_deg']),
            "{:.1f}".format(s['F_strut_kN']),
            "{:.1f}".format(s['F_tie_x_kN']),
            "{:.1f}".format(s['F_tie_y_kN']),
        ])
    _make_table(doc, ['Strut', 'L (mm)', 'θ (deg)',
                      'F_strut (kN)', 'F_tie_x (kN)', 'F_tie_y (kN)'], rows)

    # 4. Capacity Checks
    doc.add_heading('4. Capacity Checks (ACI 318-19 Ch. 23)', level=1)
    p = doc.add_paragraph()
    p.add_run('Strut effective compressive strength: ').bold = True
    p.add_run("f_ce = 0.85 x β_s x f'c (ACI 318-19 Eq. 23.4.3a) "
              "= 0.85 x 0.75 x {:.0f} = {:.2f} MPa".format(
                  inputs['fc'], results['fce_strut_MPa']))
    p = doc.add_paragraph()
    p.add_run('Strength reduction factors: ').bold = True
    p.add_run("φ_STM = 0.75 (ACI 21.2.1), φ_bearing = 0.75 (ACI 21.2.1(e))")
    if results.get('capacity_model_note'):
        p = doc.add_paragraph()
        rr = p.add_run('Preliminary capacity model: ')
        rr.bold = True
        p.add_run(results['capacity_model_note'])

    # Node type: CTT (βn=0.60) for n>=4 piles; CCT (βn=0.80) for n<4 piles
    _n_pile = results.get('n_piles', 4)
    _bn_pile = results.get('bn_pile', 0.60 if _n_pile >= 4 else 0.80)
    _node_type = "CTT" if _n_pile >= 4 else "CCT"
    _node_label = "{}, βn={:.2f}".format(_node_type, _bn_pile)

    rows = [
        ["Strut compression (φFns)",
         "{:.0f}".format(results['phi_Fns_kN']),
         "{:.0f}".format(results['F_strut_max_kN']),
         "{:.2f}".format(results['strut_DCR']),
         "OK" if results['strut_DCR'] <= 1 else "FAIL"],
        ["Pile bearing ({})".format(_node_label),
         "{:.0f}".format(results['phi_Pn_bearing_kN']),
         "{:.0f}".format(results['P_max_kN']),
         "{:.2f}".format(results['bearing_DCR']),
         "OK" if results['bearing_DCR'] <= 1 else "FAIL"],
        ["Column bearing (CCC, βn=1.00)",
         "{:.0f}".format(results['phi_Pn_column_kN']),
         "{:.0f}".format(results['Pu_kN']),
         "{:.2f}".format(results['column_DCR']),
         "OK" if results['column_DCR'] <= 1 else "FAIL"],
        ["Min strut angle ≥25° ({:.1f}°)".format(
            results['min_strut_angle_deg']),
         "-", "-", "-",
         "OK" if results['angle_OK'] else "FAIL"],
        ["Uplift / tension pile",
         "-", "{:.0f}".format(results['P_min_kN']), "-",
         "OK" if not results.get('has_uplift') else "FAIL"],
        ["Reaction equilibrium",
         "-", "-", "-",
         "OK" if results.get('reaction_equilibrium_OK', True) else "FAIL"],
    ]
    _make_table(doc, ['Check', 'Capacity (kN)',
                      'Demand (kN)', 'DCR', 'Status'], rows)

    # 5. Reinforcement
    doc.add_heading('5. Tie Reinforcement Design (ACI 318-19 §23.7)', level=1)

    doc.add_heading('5.1 Design Basis', level=2)
    p = doc.add_paragraph()
    p.add_run('Governing formula (ACI §23.7.2): ').bold = True
    p.add_run('As ≥ F_tie / (φ × fy)    where  φ = 0.75 (ACI Table 21.2.1)')
    p = doc.add_paragraph()
    p.add_run('Minimum flexural reinforcement check: ').bold = True
    p.add_run(
        'Bottom-face reinforcement in each direction is also checked against '
        'As_min = 0.0018Ag. Therefore As_req = max(As_STM, 0.0018Ag).')
    p = doc.add_paragraph()
    p.add_run('Tie force principle (STM equilibrium): ').bold = True
    p.add_run(
        'F_tie = horizontal component of the diagonal strut. '
        'For each pile, use dx_i = x_pile - x_col and '
        'dy_i = y_pile - y_col with reaction P_i:')
    p = doc.add_paragraph()
    p.add_run('  F_tie_x,i = P_i × |dx_i| / d_eff      '
              'F_tie_y,i = P_i × |dy_i| / d_eff')
    p = doc.add_paragraph()
    p.add_run(
        'The total tie force per direction is the SUM over all piles '
        'contributing to that direction (ACI STM equilibrium, not the single maximum):')
    p = doc.add_paragraph()
    p.add_run('  ΣF_tie_x = Σ (P_i × |dx_i|) / d_eff    (i = piles on controlling side)')
    p = doc.add_paragraph()
    p.add_run('  ΣF_tie_y = Σ (P_i × |dy_i|) / d_eff    (i = piles on controlling side)')

    doc.add_heading('5.2 Per-Pile Tie Force Contribution', level=2)
    p = doc.add_paragraph()
    p.add_run('d_eff = {:.0f} mm    φ = 0.75    fy_x = {:.0f} MPa    fy_y = {:.0f} MPa'.format(
        results['d_effective_mm'],
        results.get('fy_x_design_mpa', inputs['fy']),
        results.get('fy_y_design_mpa', inputs['fy']))).bold = True

    # Build per-pile table
    struts = results.get('struts', [])
    pile_rows = []
    Ftx_sum = 0.0; Fty_sum = 0.0
    for i, s in enumerate(struts):
        x_i, y_i = s['coord']
        p_i = s['P_i_kN']
        ftx = s['F_tie_x_kN']
        fty = s['F_tie_y_kN']
        dx_i = s.get('dx_from_col', x_i)
        dy_i = s.get('dy_from_col', y_i)
        Ftx_sum += ftx; Fty_sum += fty
        pile_rows.append([
            'P{}'.format(i+1),
            '{:.0f}'.format(x_i), '{:.0f}'.format(y_i),
            '{:.1f}'.format(p_i),
            '{:.0f}'.format(dx_i), '{:.0f}'.format(dy_i),
            '{:.1f}'.format(ftx), '{:.1f}'.format(fty),
        ])
    pile_rows.append([
        'TOTAL', '—', '—', '—', '—', '—',
        '{:.1f}'.format(results.get(
            'F_tie_x_design_kN',
            results['As_x_required_mm2'] * 0.75 * inputs['fy'] / 1000)),
        '{:.1f}'.format(results.get(
            'F_tie_y_design_kN',
            results['As_y_required_mm2'] * 0.75 * inputs['fy'] / 1000)),
    ])
    _make_table(doc,
        ['Pile', 'x (mm)', 'y (mm)', 'P_i (kN)',
         'dx from col (mm)', 'dy from col (mm)',
         'F_tie_x (kN)', 'F_tie_y (kN)'],
        pile_rows)

    doc.add_heading('5.3 Tie Force Summary & As Calculation', level=2)
    is_3p = results.get('is_3pile_resultant', False)
    fyx = results.get('fy_x_design_mpa', inputs['fy'])
    fyy = results.get('fy_y_design_mpa', inputs['fy'])
    ftx_design = results.get(
        'F_tie_x_design_kN',
        results['As_x_required_mm2'] * 0.75 * fyx / 1000)
    fty_design = results.get(
        'F_tie_y_design_kN',
        results['As_y_required_mm2'] * 0.75 * fyy / 1000)
    As_x_stm = results.get(
        'As_x_stm_required_mm2',
        ftx_design * 1000.0 / (0.75 * fyx) if fyx > 0 else 0.0)
    As_y_stm = results.get(
        'As_y_stm_required_mm2',
        fty_design * 1000.0 / (0.75 * fyy) if fyy > 0 else 0.0)
    As_x_min = results.get(
        'As_x_min_required_mm2',
        0.0018 * inputs['cap_ly'] * inputs['h_cap'])
    As_y_min = results.get(
        'As_y_min_required_mm2',
        0.0018 * inputs['cap_lx'] * inputs['h_cap'])
    if is_3p:
        p = doc.add_paragraph()
        p.add_run('3-Pile resultant tie: ').bold = True
        p.add_run(
            'For 3-pile caps, the resultant tie force is used to avoid '
            'coordinate-rotation dependency:')
        p = doc.add_paragraph()
        p.add_run('  F_res = √(ΣFtx² + ΣFty²) = {:.1f} kN'.format(
            results['F_tie_res_kN']))
        p = doc.add_paragraph()
        p.add_run('  As_x = F_res / (φ × fy_x) = '
                  '{:.1f} / (0.75 × {:.0f}) = {:.0f} mm²; '
                  'As_y = F_res / (φ × fy_y) = '
                  '{:.1f} / (0.75 × {:.0f}) = {:.0f} mm²'.format(
                      results['F_tie_res_kN'], fyx,
                      As_x_stm,
                      results['F_tie_res_kN'], fyy,
                      As_y_stm))
    else:
        p = doc.add_paragraph()
        p.add_run('X-direction: ').bold = True
        p.add_run(
            'ΣF_tie_x = {:.1f} kN  →  '
            'As_STM,x = {:.1f} / (0.75 × {:.0f}) = {:.0f} mm²'.format(
                ftx_design, ftx_design, fyx, As_x_stm))
        p = doc.add_paragraph()
        p.add_run('Y-direction: ').bold = True
        p.add_run(
            'ΣF_tie_y = {:.1f} kN  →  '
            'As_STM,y = {:.1f} / (0.75 × {:.0f}) = {:.0f} mm²'.format(
                fty_design, fty_design, fyy, As_y_stm))

    _make_table(doc,
        ['Direction', 'As_STM (mm²)', 'As_min = 0.0018Ag (mm²)',
         'As_req = max (mm²)', 'Governing'],
        [['X',
          '{:.0f}'.format(As_x_stm),
          '{:.0f}'.format(As_x_min),
          '{:.0f}'.format(results['As_x_required_mm2']),
          results.get('As_x_governs', '—')],
         ['Y',
          '{:.0f}'.format(As_y_stm),
          '{:.0f}'.format(As_y_min),
          '{:.0f}'.format(results['As_y_required_mm2']),
          results.get('As_y_governs', '—')]])

    doc.add_heading('5.4 Selected Reinforcement', level=2)
    _sx_spacing = _rebar_spacing(
        x_chk['n_bars'], inputs['cap_ly'], inputs['cap_lx'],
        inputs['cap_ly'], inputs['cover'])
    _sy_spacing = _rebar_spacing(
        y_chk['n_bars'], inputs['cap_lx'], inputs['cap_lx'],
        inputs['cap_ly'], inputs['cover'])
    _make_table(doc,
        ['Direction', 'ΣF_tie (kN)', 'As req (mm²)', 'Governing',
         'fy used (MPa)', 'Selected', 'Spacing s (mm)',
         'As prov (mm²)', 'Min OK', 'STM OK', 'Ratio', 'Status'],
        [["X",
          "{:.1f}".format(results['F_tie_x_max_kN'] if not is_3p else results['F_tie_res_kN']),
          "{:.0f}".format(results['As_x_required_mm2']),
          results.get('As_x_governs', '—'),
          "{:.0f}".format(fyx),
          "{}-{}".format(x_chk['n_bars'], x_chk['bar_size']),
          _format_spacing(_sx_spacing),
          "{:.0f}".format(x_chk['As_provided']),
          "OK" if x_chk['As_provided'] >= As_x_min else "FAIL",
          "OK" if x_chk.get('force_ok', True) else "FAIL",
          "{:.2f}".format(x_chk['ratio']),
          "OK" if x_chk['ok'] else "FAIL"],
         ["Y",
          "{:.1f}".format(results['F_tie_y_max_kN'] if not is_3p else results['F_tie_res_kN']),
          "{:.0f}".format(results['As_y_required_mm2']),
          results.get('As_y_governs', '—'),
          "{:.0f}".format(fyy),
          "{}-{}".format(y_chk['n_bars'], y_chk['bar_size']),
          _format_spacing(_sy_spacing),
          "{:.0f}".format(y_chk['As_provided']),
          "OK" if y_chk['As_provided'] >= As_y_min else "FAIL",
          "OK" if y_chk.get('force_ok', True) else "FAIL",
          "{:.2f}".format(y_chk['ratio']),
          "OK" if y_chk['ok'] else "FAIL"]])

    # Figure: Bottom Rebar Layout
    img_bot = _plot_bottom_rebar_fig(
        inputs['coords'], inputs['D'],
        inputs['cap_lx'], inputs['cap_ly'],
        inputs.get('cap_cx', 0), inputs.get('cap_cy', 0),
        inputs['col_size'], inputs.get('cap_polygon'),
        x_chk['bar_size'], x_chk['n_bars'],
        y_chk['bar_size'], y_chk['n_bars'],
        x_chk, y_chk, inputs['cover'])
    if img_bot:
        doc.add_picture(img_bot, width=Inches(6))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run(
            'Figure 3: Bottom-Face Reinforcement Layout — '
            '{} × {} (X-dir, red) and {} × {} (Y-dir, blue dashed)'.format(
                x_chk['n_bars'], x_chk['bar_size'],
                y_chk['n_bars'], y_chk['bar_size'])).italic = True
    
    # 5.5 Auto-Optimized Rebar (min weight)
    if opt_x and opt_y:
        doc.add_heading('5.5 Auto-Optimized Rebar (Minimum Weight)', level=2)
        p = doc.add_paragraph()
        p.add_run('Optimization criterion: ').bold = True
        p.add_run("Minimum total steel weight while As_provided ≥ As_required, "
                  "bar lengths assumed equal to cap span in each direction.")
        _make_table(doc,
            ['Direction', 'Optimal selection', 'As prov (mm²)',
             'Bar length (mm)', 'Weight (kg)', 'Status'],
            [["X (along x-axis)",
              "{}-{}".format(opt_x["n_bars"], opt_x["bar_size"]),
              "{:.0f}".format(opt_x["As_provided"]),
              "{:.0f}".format(opt_x["bar_length_mm"]),
              "{:.2f}".format(opt_x["weight_kg"]),
              "OK" if opt_x["ok"] else "FAIL"],
             ["Y (along y-axis)",
              "{}-{}".format(opt_y["n_bars"], opt_y["bar_size"]),
              "{:.0f}".format(opt_y["As_provided"]),
              "{:.0f}".format(opt_y["bar_length_mm"]),
              "{:.2f}".format(opt_y["weight_kg"]),
              "OK" if opt_y["ok"] else "FAIL"]])

    # 5.6 Anchorage / Development Length (ACI 25.4)
    if anch_x and anch_y:
        doc.add_heading('5.6 Anchorage at CCT Nodes (ACI 318-19 §25.4)', level=2)
        p = doc.add_paragraph()
        p.add_run('Per ACI 23.8.3, ').bold = True
        p.add_run("tie reinforcement shall develop fy at the point where the "
                  "centroid of the tie crosses the extended nodal zone. "
                  "Required ld and ldh are checked below:")
        p = doc.add_paragraph()
        p.add_run("• ld (straight)  ≈ (fy·ψs / 1.1λ√f'c·(cb+Ktr)/db)·db   "
                  "(ACI Eq. 25.4.2.3a)")
        p = doc.add_paragraph()
        p.add_run("• ldh (hook)     ≈ (fy / 23λ√f'c)·db^1.5             "
                  "(ACI Eq. 25.4.3.1a)")
        _make_table(doc,
            ['Direction', 'Bar', 'ld req (mm)', 'ldh req (mm)',
             'Avail straight (mm)', 'Avail hook (mm)', 'Recommended'],
            [["X",
              anch_x["bar_size"],
              "{:.0f}".format(anch_x["ld_required_mm"]),
              "{:.0f}".format(anch_x["ldh_required_mm"]),
              "{:.0f}".format(anch_x["available_straight_mm"]),
              "{:.0f}".format(anch_x["available_hook_mm"]),
              anch_x["recommended"]],
             ["Y",
              anch_y["bar_size"],
              "{:.0f}".format(anch_y["ld_required_mm"]),
              "{:.0f}".format(anch_y["ldh_required_mm"]),
              "{:.0f}".format(anch_y["available_straight_mm"]),
              "{:.0f}".format(anch_y["available_hook_mm"]),
              anch_y["recommended"]]])

    # 6. Top-Face Minimum Reinforcement
    doc.add_heading('6. Top-Face Minimum Reinforcement (ACI 318-19)', level=1)
    if top_rebar:
        tr = top_rebar
        p = doc.add_paragraph()
        p.add_run(
            "The top face of the pile cap is not part of the STM load path "
            "and is designed here using one-half of the gross-area minimum "
            "reinforcement: As_top = (0.0018Ag)/2 in each direction."
        )

        doc.add_heading('6.1 Top-Face Minimum — (0.0018Ag)/2',
                        level=2)
        p = doc.add_paragraph()
        p.add_run("ρ_top = 0.0018 / 2 = {:.4f}".format(
            tr.get("rho_top", 0.0009))).bold = True
        _make_table(doc, ['Direction', 'b (mm)', 'Ag = b×h (mm²)',
                          '0.0018Ag (mm²)', 'As_top = (0.0018Ag)/2 (mm²)'], [
            ['X', '{:.0f}'.format(inputs['cap_ly']),
             '{:.0f}'.format(tr.get('Ag_x_mm2', inputs['cap_ly'] * inputs['h_cap'])),
             '{:.0f}'.format(tr.get('As_full_min_x_mm2', tr.get('As_ts_x_mm2', 0))),
             '{:.0f}'.format(tr['As_top_x_mm2'])],
            ['Y', '{:.0f}'.format(inputs['cap_lx']),
             '{:.0f}'.format(tr.get('Ag_y_mm2', inputs['cap_lx'] * inputs['h_cap'])),
             '{:.0f}'.format(tr.get('As_full_min_y_mm2', tr.get('As_ts_y_mm2', 0))),
             '{:.0f}'.format(tr['As_top_y_mm2'])],
        ])

        doc.add_heading('6.2 Governing Top-Face Reinforcement', level=2)
        _s_gov = tr.get("s_max_top_mm", min(3*inputs['h_cap'], 450))
        _s_ts  = tr.get("s_ts_max_mm",  min(3*inputs['h_cap'], 450))
        _s_cr  = tr.get("s_crack_mm",   9999)
        _fs    = tr.get("fs_service_mpa", (2/3)*tr.get("fy_design_mpa", 420))
        _fy_d  = tr.get("fy_design_mpa", 420)

        p = doc.add_paragraph()
        p.add_run("Bar selected: ").bold = True
        p.add_run("{} (fy_bar = {:.0f} MPa, fy_design = {:.0f} MPa — {})".format(
            tr.get("top_bar_size", "—"),
            tr.get("fy_bar_mpa", _fy_d), _fy_d,
            tr.get("fy_note", "")))

        _make_table(doc,
                    ['Direction', 'As_top req (mm²)', 'Recommended bars',
                     'Spacing s (mm)', 'As prov (mm²)', 'Design Basis'],
                    [['X', '{:.0f}'.format(tr['As_top_x_mm2']),
                      '{}-{}'.format(tr.get('top_x_n_bars', 2),
                                     tr.get('top_bar_size', 'DB20')),
                      '{:.0f}'.format(tr.get('top_x_spacing_mm', 0.0)),
                      '{:.0f}'.format(tr.get('top_x_As_provided_mm2', 0.0)),
                      tr['governs_x']],
                     ['Y', '{:.0f}'.format(tr['As_top_y_mm2']),
                      '{}-{}'.format(tr.get('top_y_n_bars', 2),
                                     tr.get('top_bar_size', 'DB20')),
                      '{:.0f}'.format(tr.get('top_y_spacing_mm', 0.0)),
                      '{:.0f}'.format(tr.get('top_y_As_provided_mm2', 0.0)),
                      tr['governs_y']]])
        p = doc.add_paragraph()
        p.add_run("Note: ").bold = True
        p.add_run(tr.get("top_design_note", "Top-face As is governed by the minimum top mat."))

        doc.add_heading('6.3 Spacing Limits', level=2)
        _make_table(doc,
                    ['Check', 's_max (mm)', 'Active'],
                    [['§24.4.3.3  min(3h={:.0f}, 450)'.format(3*inputs['h_cap']),
                      '{:.0f}'.format(_s_ts), 'Always'],
                     ['§24.3.2  crack-width  (fs={:.0f} MPa)'.format(_fs),
                      '{:.0f}'.format(_s_cr),
                      'fy > 420' if _fy_d > 420 else 'Not active (fy ≤ 420)'],
                     ['Governing s_max', '{:.0f}'.format(_s_gov), '← use this']])

        p = doc.add_paragraph()
        p.add_run("Placement: ").bold = True
        p.add_run(
            "Top bars placed in both X and Y directions, "
            "uniformly distributed across the full cap width, "
            "at cover depth from the top surface. "
            "fy used in design per ACI §20.2.2.4: "
            "≤DB28 → 390 MPa, DB32 → 490 MPa, capped at 550 MPa.")

        # Figure: Top Rebar Layout
        img_top = _plot_top_rebar_fig(
            inputs['coords'], inputs['D'],
            inputs['cap_lx'], inputs['cap_ly'],
            inputs.get('cap_cx', 0), inputs.get('cap_cy', 0),
            inputs['col_size'], inputs.get('cap_polygon'),
            tr, inputs['cover'])
        if img_top:
            doc.add_picture(img_top, width=Inches(6))
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(
                'Figure 4: Top-Face Minimum Reinforcement Layout — '
                '{n_x} × {bar} (X-dir, amber) and '
                '{n_y} × {bar} (Y-dir, teal dashed). '
                'fy_design = {fy:.0f} MPa, s_max = {s:.0f} mm'.format(
                    n_x=tr.get('top_x_n_bars', 2),
                    n_y=tr.get('top_y_n_bars', 2),
                    bar=tr.get('top_bar_size', 'DB20'),
                    fy=tr.get('fy_design_mpa', 390),
                    s=tr.get('s_max_top_mm', 450))).italic = True
    else:
        doc.add_paragraph("Top reinforcement data not available.")

    # 7. Notes
    doc.add_heading('7. Important Notes', level=1)
    p = doc.add_paragraph()
    rr = p.add_run("Per ACI 318-19 §13.4.6.3: ")
    rr.bold = True
    p.add_run(
        "When the Strut-and-Tie Method (STM) is used to design a pile cap "
        "in accordance with ACI 318-19 Chapter 23, separate beam-shear and "
        "two-way (punching) shear checks are NOT required. The shear "
        "behavior is implicitly captured through the strength of struts "
        "and nodal zones in the STM model.")

    # 8. Conclusion
    doc.add_heading('8. Conclusion', level=1)
    p = doc.add_paragraph()
    if design_ok:
        rr = p.add_run("DESIGN OK. ")
        rr.bold = True
        rr.font.color.rgb = OK_GREEN
        p.add_run("STM capacity checks, bottom reinforcement, top minimum "
                  "reinforcement, and available anchorage checks are satisfied "
                  "for the stated input assumptions.")
    else:
        rr = p.add_run("DESIGN FAILS. ")
        rr.bold = True
        rr.font.color.rgb = FAIL_RED
        p.add_run("One or more checks are unsatisfied. Review the executive "
                  "summary and detailed tables; typical adjustments include "
                  "increasing cap thickness, pile size, cap dimensions, pile "
                  "count, reinforcement, or available anchorage length.")

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    return out
